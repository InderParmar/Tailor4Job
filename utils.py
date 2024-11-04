from groq import Groq
import requests
import json
from docx import Document 

import pdfkit

def process_files(files, api_key, model, mode, token_usage, provider_name):
    """
    Process input files using either the Groq API or OpenRouter API, depending on the provider.
    """ 
    content = ''
    
    # Read content from input files
    for file_path in files:
        if file_path.endswith('.docx'):
            # Read .docx file content
            doc = Document(file_path)
            for para in doc.paragraphs:
                content += para.text + '\n'
        else:
            # Read other files assuming they are text-based
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content += file.read() + '\n'

    # Build the prompt based on the mode (detailed or basic)
    if mode == 'detailed':
        prompt = """
        Attached are the Resume, Cover Letter, and Job Description for which the candidate is applying for the job.

        What I want from you is to analyze the Resume and Cover Letter and compare them to the Job Description.

        1. **Brief Introduction**:
           - Introduce the job and the candidate (mention the name of the candidate from the Resume).

        2. **Analysis Results**:
           - Compare the Resume to the Job Description:
             - Is the candidate a good fit or not?
             - What are the candidate's strong points for this job?
             - What are the candidate's weaknesses?
             - How can the candidate improve the Resume?
           - Compare the Cover Letter to the Job Description:
             - Is the cover letter aligned with the job requirements?
             - What improvements can be made?

        3. **Summary**:
           - Estimate the percentage chance that this Resume and Cover Letter can pass the ATS system for this job.
           - List any important keywords that can be added to the Resume and Cover Letter.
           - Suggest keywords that should be replaced with better alternatives.
        """
    else:
        prompt = """
        Attached are the Resume, Cover Letter, and Job Description for which the candidate is applying for the job.

        Please provide:
        1. A brief introduction to the job and candidate.
        2. An estimated percentage chance of the resume and cover letter passing an ATS system.
        """
    content = prompt + "\n" + content
    # Process content using the appropriate provider
    if provider_name == 'groq':
        # Initialize Groq client
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": content,
                }
            ],
            model=model,
        )
        tailored_content = response.choices[0].message.content
        token_info = response.usage if token_usage else None

    elif provider_name == 'openrouter':
        # Make OpenRouter API request
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
            },
            data=json.dumps({
                "model": model,
                "messages": [
                    {
                        "role": "user",
                        "content": content
                    }
                ]
            })
        )
        if response.status_code != 200:
            raise Exception(f"OpenRouter API request failed with status code {response.status_code}: {response.text}")
        result = response.json()
        tailored_content = result['choices'][0]['message']['content']
        token_info = result.get('usage', None) if token_usage else None

    return tailored_content, token_info

def generate_output(content, output_file):
    """
    Generate output file in either .docx or .pdf format.
    """
    if output_file.endswith('.docx'):
        generate_docx(content, output_file)
    elif output_file.endswith('.pdf'):
        generate_pdf(content, output_file)
    else:
        raise Exception('Unsupported file format. Please use .docx or .pdf.')

def generate_docx(content, output_file):
    """
    Generate a .docx file with the given content.
    """
    doc = Document()
    doc.add_paragraph(content)
    doc.save(output_file)

def generate_pdf(content, output_file):
    """
    Generate a .pdf file with the given content.
    """
    # Create a temporary HTML file to convert to PDF
    html_content = f'<html><body><p>{content}</p></body></html>'
    pdfkit.from_string(html_content, output_file)